from docx import Document
import subprocess
import os
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE

FIO = "Михайлов Евгений Игоревич"
course = "1"
faculty = "OIKS"

# style = document.styles['Normal']
# font = style.font
# font.name = 'Arial'
# font.size = Pt(10)
 




def income_statement(id:str, **kwargs ):
    document = Document()
    p = document.add_paragraph(
        f'МИНИСТЕРСТВО НАУКИ ВЫСШЕГО ОБРАЗОВАНИЯ\nРОССИЙСКОЙ ФЕДЕРАЦИИ\nФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ ОБРАЗОВАНИЕ\nУЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ\n\"Национальный исследовательский\nядерный университет \"МИФИ\"\n(ИАТЭ НИЯУ МИФИ)\nдата:  {kwargs["date"]}\n\n\n\n\n\n\n\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('СПРАВКА О ДОХОДАХ').alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph('').alignment = WD_ALIGN_PARAGRAPH.CENTER

    font_styles = document.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(16)
    font_object.name = 'Times New Roman'

    table = document.add_table(rows=2 + len(kwargs['stlist'])+1, cols=4)
    table.cell(0, 0).merge(table.cell(0, 3))
    cell_name = table.cell(0, 0)
    cell_name.text = kwargs['FIO']

    cell_overall = table.cell(1, 3)
    cell_overall.text = 'Всего'

    cell_month = table.cell(1, 0)
    cell_month.text = 'Месяц'

    table_academic_scholarship = table.cell(1, 1)
    table_academic_scholarship.text = 'Академическая стипендия'

    table_social_scholarship = table.cell(1, 2)
    table_social_scholarship.text = 'Социальная стипендия'
    alot = 0
    for i, item in enumerate(kwargs['stlist']):
        count = 0
        table.cell(2 + i, 0).text = item[0]
        if id in item[1]:
            table.cell(2 + i, 1).text = "1800"
            count +=1800
        if id in item[2]:
            table.cell(2 + i, 2).text = "1300"
            count +=1300
        alot+=count

        table.cell(2 + i, 3).text = str(count)
    table.cell(2+len(kwargs['stlist']),3).text = str(alot)

    document.save('demo.docx')

    pic = document.add_picture('C:/DjangoProjects/hackatonSite/main/static/images/print.jpg', width=Inches(1.25))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save('C:/DjangoProjects/hackatonSite/main/static/images/demo.docx')


