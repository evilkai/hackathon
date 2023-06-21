from docx import Document
import subprocess
import os
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE

FIO = "Михайлов Евгений Игоревич"
course = "1"
faculty = "OIKS"

# style = document.styles['Normal']
# font = style.font
# font.name = 'Arial'
# font.size = Pt(10)
 




def study_statement(**kwargs):
    
    document = Document()


    p = document.add_paragraph(
        f'МИНИСТЕРСТВО НАУКИ ВЫСШЕГО ОБРАЗОВАНИЯ\nРОССИЙСКОЙ ФЕДЕРАЦИИ\nФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ ОБРАЗОВАНИЕ\nУЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ\n\"Национальный исследовательский\nядерный университет \"МИФИ\"\n(ИАТЭ НИЯУ МИФИ)\nдата: {kwargs["date"]}\n\n\n\n\n\n\n\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('СПРАВКА').alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph('').alignment = WD_ALIGN_PARAGRAPH.CENTER

    font_styles = document.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(14)
    font_object.name = 'Times New Roman'

    parag = document.add_paragraph("")

    parag.add_run(
        f'           О том, что {kwargs["FIO"]} является студентом(кой) {kwargs["course"]} курса ({kwargs["form"]} форма обучения) направления {kwargs["direction"]} Обнинского института атомной энергетики - филиала федерального государственного автономного образовательного учреждения высшего образования "Национальный исследовательский ядерный университет "МИФИ" (ИАТЭ НИЯУ МИФИ).\n',
        style='CommentsStyle')

    parag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parag_last = document.add_paragraph("")
    parag_last.add_run(
        f'           Гражданство: {kwargs["country"]}\n'
        f'           Основания обучения: {kwargs["basis"]}\n'
        f'           Справка дана для предъявления по месту требования.\n',

    style='CommentsStyle')
    parag_last.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pic = document.add_picture('C:/DjangoProjects/hackatonSite/main/static/images/print.jpg', width=Inches(1.25))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save('C:/DjangoProjects/hackatonSite/main/static/images/demo.docx')


